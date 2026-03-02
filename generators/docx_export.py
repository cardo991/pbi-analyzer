"""Data Dictionary export as Word (.docx) document."""

import io
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def generate_data_dictionary(project_name: str, documentation: dict, score: dict,
                             findings: list, branding: dict = None, lang: str = "en") -> io.BytesIO:
    """Generate a Word document data dictionary.

    Returns BytesIO with the .docx content.
    Raises ImportError if python-docx is not installed.
    """
    if not HAS_DOCX:
        raise ImportError("python-docx is required. Run: pip install python-docx")

    doc = Document()

    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Cover Page ---
    _add_cover(doc, project_name, score, branding)
    doc.add_page_break()

    # --- Table of Contents placeholder ---
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('(Update this field in Word to generate TOC)')
    doc.add_page_break()

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    overview = documentation.get('overview', {})
    p = doc.add_paragraph()
    p.add_run(f'Project: ').bold = True
    p.add_run(project_name)
    p = doc.add_paragraph()
    p.add_run(f'Score: ').bold = True
    p.add_run(f"{score.get('total_score', 0)}/100 (Grade: {score.get('grade', 'F')})")
    p = doc.add_paragraph()
    p.add_run(f'Analyzed: ').bold = True
    p.add_run(datetime.now().strftime('%Y-%m-%d %H:%M'))

    # Category scores
    cat_scores = score.get('category_scores', {})
    if cat_scores:
        doc.add_heading('Category Scores', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = 'Category'
        hdr[1].text = 'Score'
        for cat, val in cat_scores.items():
            row = table.add_row().cells
            row[0].text = cat.replace('_', ' ').title()
            row[1].text = str(round(val, 1)) if isinstance(val, (int, float)) else str(val)

    doc.add_page_break()

    # --- Model Overview ---
    doc.add_heading('Model Overview', level=1)
    stats = [
        ('Total Tables', overview.get('total_tables', 0)),
        ('Total Measures', overview.get('total_measures', 0)),
        ('Total Columns', overview.get('total_columns', 0)),
        ('Total Relationships', overview.get('total_relationships', 0)),
        ('Data Sources', overview.get('data_sources_count', 0)),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    for name, val in stats:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = str(val)

    doc.add_page_break()

    # --- Tables & Columns ---
    tables_data = documentation.get('tables', [])
    if tables_data:
        doc.add_heading('Tables & Columns', level=1)
        for tbl in tables_data:
            tbl_name = tbl.get('name', 'Unknown')
            doc.add_heading(tbl_name, level=2)

            # Table info
            row_count = tbl.get('row_count', '')
            if row_count:
                doc.add_paragraph(f'Rows: {row_count}')

            # Columns
            columns = tbl.get('columns', [])
            if columns:
                doc.add_heading('Columns', level=3)
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Grid Accent 1'
                hdr = table.rows[0].cells
                hdr[0].text = 'Column'
                hdr[1].text = 'Data Type'
                hdr[2].text = 'Description'
                for col in columns:
                    row = table.add_row().cells
                    row[0].text = col.get('name', '')
                    row[1].text = col.get('data_type', '')
                    row[2].text = col.get('description', '')

    doc.add_page_break()

    # --- Measure Dictionary ---
    measures_data = documentation.get('measures', [])
    if measures_data:
        doc.add_heading('Measure Dictionary', level=1)

        # Group by table
        by_table = {}
        for m in measures_data:
            tbl = m.get('table', 'Unknown')
            by_table.setdefault(tbl, []).append(m)

        for tbl_name, measures in sorted(by_table.items()):
            doc.add_heading(tbl_name, level=2)
            for m in measures:
                p = doc.add_paragraph()
                run = p.add_run(m.get('name', ''))
                run.bold = True
                run.font.size = Pt(11)

                if m.get('expression'):
                    expr_p = doc.add_paragraph()
                    expr_run = expr_p.add_run(m['expression'])
                    expr_run.font.name = 'Consolas'
                    expr_run.font.size = Pt(9)
                    expr_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                if m.get('format_string'):
                    doc.add_paragraph(f"Format: {m['format_string']}", style='List Bullet')

                if m.get('description'):
                    doc.add_paragraph(m['description'])

                doc.add_paragraph('')  # spacer

        doc.add_page_break()

    # --- Relationships ---
    relationships = documentation.get('relationships', [])
    if relationships:
        doc.add_heading('Relationships', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'From'
        hdr[1].text = 'To'
        hdr[2].text = 'Cardinality'
        hdr[3].text = 'Cross-Filter'
        for r in relationships:
            row = table.add_row().cells
            row[0].text = f"{r.get('from_table', '')}[{r.get('from_column', '')}]"
            row[1].text = f"{r.get('to_table', '')}[{r.get('to_column', '')}]"
            row[2].text = f"{r.get('from_cardinality', '')}:{r.get('to_cardinality', '')}"
            row[3].text = r.get('cross_filtering', '')

        doc.add_page_break()

    # --- Data Sources ---
    sources = documentation.get('data_sources', [])
    if sources:
        doc.add_heading('Data Sources', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Table'
        hdr[1].text = 'Source Type'
        hdr[2].text = 'Details'
        for s in sources:
            row = table.add_row().cells
            row[0].text = s.get('table', '')
            row[1].text = s.get('source_type', '')
            row[2].text = s.get('detail', '')

        doc.add_page_break()

    # --- Findings Summary ---
    if findings:
        doc.add_heading('Findings Summary', level=1)

        # Group by severity
        by_severity = {'error': [], 'warning': [], 'info': []}
        for f in findings:
            sev = f.get('severity', 'info') if isinstance(f, dict) else getattr(f, 'severity', 'info')
            by_severity.setdefault(sev, []).append(f)

        for sev in ('error', 'warning', 'info'):
            items = by_severity.get(sev, [])
            if not items:
                continue
            doc.add_heading(f'{sev.title()} ({len(items)})', level=2)
            for f in items[:20]:  # Limit to 20 per severity
                msg = f.get('message', '') if isinstance(f, dict) else getattr(f, 'message', '')
                loc = f.get('location', '') if isinstance(f, dict) else getattr(f, 'location', '')
                rule = f.get('rule_id', '') if isinstance(f, dict) else getattr(f, 'rule_id', '')
                p = doc.add_paragraph(style='List Bullet')
                if rule:
                    run = p.add_run(f'[{rule}] ')
                    run.bold = True
                p.add_run(msg)
                if loc:
                    p.add_run(f' — {loc}').italic = True

    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _add_cover(doc, project_name, score, branding):
    """Add a cover page."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph('')

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DATA DICTIONARY')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x77, 0xCC)
    run.bold = True

    # Project name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project_name)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Score
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Score: {score.get('total_score', 0)}/100 — Grade: {score.get('grade', 'F')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime('%B %d, %Y'))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Branding
    if branding and branding.get('company_name'):
        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(branding['company_name'])
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0x77, 0xCC)

    # GridPulse attribution
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Generated by GridPulse — Power BI Intelligence Platform')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
